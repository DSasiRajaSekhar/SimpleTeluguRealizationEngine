from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Pt, RGBColor


PAPER = Path(r"D:\NLG\Paper.docx")
BACKUP = Path(r"D:\NLG\Paper_before_reference_update.docx")


RELATED_WORK_PARAGRAPH = (
    "The broader multilingual SimpleNLG literature also supports this view of the Telugu system. "
    "Adaptations of SimpleNLG to Italian, Spanish, and Mandarin show that the general architecture "
    "of a realizer can be reused across languages only when language-specific morphology, agreement, "
    "and ordering constraints are made explicit (Mazzei, Battaglino, and Bosco, 2016; Ramos-Soto, "
    "Janeiro-Gallardo, and Bugarin, 2017; Chen, van Deemter, and Lin, 2018). These later adaptations "
    "are useful comparators for the present work because they treat surface realization as a modular "
    "engineering problem rather than as a purely language-independent mapping problem. Recent work on "
    "low-resource surface-realization architecture makes a similar point: for under-resourced languages, "
    "inspectable rules, reusable resources, and reproducible evaluation are central to extending a "
    "realizer beyond a small demonstration system (Mahlaza and Keet, 2023)."
)


FUTURE_WORK_PARAGRAPHS = [
    (
        "The present paper remains deliberately conservative: it stabilizes the existing XML-driven "
        "rule-based realization pipeline and adds two controlled extensions, rather than replacing the "
        "system with a new representation or a neural generator. A natural next step is to connect the "
        "XML input layer to a more standard syntactic representation. The Telugu dependency treebank "
        "described by Rama and Vajjala (2018) provides a useful basis for such a conversion, because it "
        "would allow future experiments to compare realization from the current hand-authored XML inputs "
        "with realization from dependency-structured inputs."
    ),
    (
        "Other recent Telugu and Dravidian resources can support later extensions but are not part of "
        "the present implementation. The Hindi-Telugu parallel corpus can support downstream evaluation "
        "in translation or controlled generation settings (Mujadia and Sharma, 2022). A Transformer-based "
        "Telugu morphological analyzer and recent Tamil-Telugu sandhi-splitting work can be used later "
        "as comparison points or backoff components for forms that are outside the present rule inventory "
        "(Dasari et al., 2023; Dasari et al., 2025). These resources are therefore best treated as future "
        "evaluation and extension paths, while the current paper reports only the rule-based changes that "
        "have been implemented and regression-tested."
    ),
]


REFERENCES = [
    "Chen, Y., van Deemter, K., and Lin, C. (2018). SimpleNLG-ZH: A linguistic realisation engine for Mandarin. INLG 2018. https://aclanthology.org/W18-6506/",
    "Dasari, H., et al. (2023). Transformer-based context aware morphological analyzer for Telugu. DravidianLangTech 2023. https://aclanthology.org/2023.dravidianlangtech-1.4/",
    "Dasari, H., et al. (2025). Sandhi splitting in Tamil and Telugu: A sequence-to-sequence approach leveraging Transformer models. CHiPSAL 2025. https://aclanthology.org/2025.chipsal-1.9/",
    "Dokkara, S. R. S., Penumathsa, S., and Sripada, S. G. (2015). A simple surface realization engine for Telugu. ENLG 2015, 1-8. https://aclanthology.org/W15-4701/",
    "Dokkara, S. R. S., Penumathsa, S., and Sripada, S. G. (2017). Morphological generator for Telugu nouns and pronouns. International Journal of Computer Applications, 165(5), 6-14. DOI: 10.5120/ijca2017913857.",
    "Dokkara, S. R. S., Penumathsa, S., and Sripada, S. G. (2017). Verb morphological generator for Telugu. Indian Journal of Science and Technology, 10(13), 1-11. DOI: 10.17485/ijst/2017/v10i13/110448.",
    "Gatt, A., and Reiter, E. (2009). SimpleNLG: A realisation engine for practical applications. ENLG 2009. https://aclanthology.org/W09-0613/",
    "Krishnamurti, B., and Gwynn, J. P. L. A Grammar of Modern Telugu. Oxford University Press.",
    "Mahlaza, Z., and Keet, C. M. (2023). Surface realization architecture for low-resourced African languages. ACM Transactions on Asian and Low-Resource Language Information Processing. DOI: 10.1145/3567594.",
    "Mazzei, A., Battaglino, C., and Bosco, C. (2016). SimpleNLG-IT: Adapting SimpleNLG to Italian. INLG 2016. DOI: 10.18653/v1/W16-6630.",
    "Mujadia, V., and Sharma, D. M. (2022). The LTRC Hindi-Telugu parallel corpus. LREC 2022. https://aclanthology.org/2022.lrec-1.365/",
    "Rama, T., and Vajjala, S. (2018). A dependency treebank for Telugu. TLT 2018. https://aclanthology.org/W17-7616.pdf",
    "Ramos-Soto, A., Janeiro-Gallardo, J., and Bugarin, A. (2017). Adapting SimpleNLG to Spanish. INLG 2017. DOI: 10.18653/v1/W17-3521.",
]


def set_run_format(run, bold=False):
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold


def add_paragraph_after(paragraph, text, style_name="Normal"):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style_name
    run = new_para.add_run(text)
    set_run_format(run, bold=style_name.startswith("Heading"))
    return new_para


def add_paragraph_end(doc, text, style_name="Normal"):
    paragraph = doc.add_paragraph(style=style_name)
    run = paragraph.add_run(text)
    set_run_format(run, bold=style_name.startswith("Heading"))
    return paragraph


def has_text(doc, text):
    return any(paragraph.text.strip() == text for paragraph in doc.paragraphs)


def main():
    if not BACKUP.exists():
        shutil.copy2(PAPER, BACKUP)

    doc = Document(PAPER)

    if RELATED_WORK_PARAGRAPH not in [p.text for p in doc.paragraphs]:
        anchor = next(
            p for p in doc.paragraphs
            if p.text.strip().startswith("Surface realization has often been studied")
        )
        add_paragraph_after(anchor, RELATED_WORK_PARAGRAPH, "Normal")

    if not has_text(doc, "Future Extensions"):
        eval_anchor = next(
            p for p in doc.paragraphs
            if p.text.strip().startswith("The coupled proper-noun extension was evaluated")
        )
        last = add_paragraph_after(eval_anchor, "Future Extensions", "Heading 1")
        for text in FUTURE_WORK_PARAGRAPHS:
            last = add_paragraph_after(last, text, "Normal")

    if not has_text(doc, "References"):
        add_paragraph_end(doc, "References", "Heading 1")
        for ref in REFERENCES:
            add_paragraph_end(doc, ref, "Normal")

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text:
                set_run_format(run, bold=paragraph.style.name.startswith("Heading"))

    doc.save(PAPER)
    print(f"Updated {PAPER}")
    print(f"Backup {BACKUP}")


if __name__ == "__main__":
    main()
