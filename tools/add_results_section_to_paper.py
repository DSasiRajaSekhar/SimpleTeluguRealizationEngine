from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


PAPER = Path(r"D:\NLG\Paper.docx")
BACKUP = Path(r"D:\NLG\Paper_before_results_section.docx")


def insert_after(paragraph, text=""):
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.clear()
    if text:
        inserted.add_run(text)
    return inserted


def format_paragraph(paragraph, bold=False):
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = bold


def main():
    if not PAPER.exists():
        raise FileNotFoundError(PAPER)

    if not BACKUP.exists():
        shutil.copy2(PAPER, BACKUP)

    doc = Document(PAPER)
    heading_text = "Evaluation of the Quantifier Extension"
    if any(p.text.strip() == heading_text for p in doc.paragraphs):
        print("Section already present; no changes made.")
        return

    anchor = doc.paragraphs[-1]
    heading = insert_after(anchor, heading_text)
    format_paragraph(heading, bold=True)

    p1_text = (
        "The numeral-quantifier extension was evaluated using the newly added "
        "Krishnamurti-based examples for person-count and object-count expressions. "
        "The system now generates ixxaru prakAsAlu from a representation containing "
        "reMdu as a person-count quantifier and prakAsaM as the plural head noun. "
        "It also generates mugguru rAmamurwulu from mUdu as a person-count quantifier, "
        "showing that compact person-count forms are produced by the quantifier "
        "mechanism rather than by placing the surface form directly in the input XML."
    )
    p1 = insert_after(heading, p1_text)
    format_paragraph(p1)

    p2_text = (
        "The same evaluation also confirms that the extension does not overgenerate "
        "person-count forms in non-person contexts. The object-count example reMdu "
        "pustakAlu is realized with the numeral unchanged, while the larger "
        "person-count expression iravE Exu is realized as iravE Exu maMxi through "
        "the fallback rule. The original fourteen evaluation XML files were run "
        "again after this change and their outputs remained unchanged. The "
        "Krishnamurti evaluation batch also continued to produce the earlier "
        "non-count outputs, with only the deliberately updated count-noun examples "
        "showing the new quantifier behavior."
    )
    p2 = insert_after(p1, p2_text)
    format_paragraph(p2)

    doc.save(PAPER)
    print(f"Updated {PAPER}")
    print(f"Backup {BACKUP}")


if __name__ == "__main__":
    main()
