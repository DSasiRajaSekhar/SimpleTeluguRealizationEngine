from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.text.paragraph import Paragraph


PAPER = Path(r"D:\NLG\Paper.docx")
BACKUP = Path(r"D:\NLG\Paper_before_quantifier_section.docx")


def insert_after(paragraph, text="", style=None):
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.clear()
    if text:
        inserted.add_run(text)
    if style:
        inserted.style = style
    return inserted


def main():
    if not PAPER.exists():
        raise FileNotFoundError(PAPER)

    if not BACKUP.exists():
        shutil.copy2(PAPER, BACKUP)

    doc = Document(PAPER)

    marker = "A central strength of the original system is that it treats Telugu surface realization as a morphology-sensitive problem."
    insert_index = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith(marker):
            insert_index = i
            break

    if insert_index is None:
        raise RuntimeError("Could not find insertion point.")

    heading_text = "Extension for Numeral Quantifiers and Person-Count Expressions"
    if any(p.text.strip() == heading_text for p in doc.paragraphs):
        print("Section already present; no changes made.")
        return

    anchor = doc.paragraphs[insert_index]
    heading = insert_after(anchor, heading_text, "Heading 1")

    p1 = (
        "A further extension concerns the treatment of numeral quantifiers, especially "
        "person-count expressions. In Telugu, numerals used with human referents may "
        "surface differently from the same numerals used with non-human or object nouns. "
        "For example, the numeral reMdu remains unchanged in an object-count expression "
        "such as reMdu pustakAlu, but the corresponding person-count expression is "
        "realized as ixxaru. Similarly, mUdu is realized as mugguru in person-count "
        "contexts. Treating these numerals as ordinary nouns and applying regular plural "
        "formation leads to incorrect outputs. Therefore, the extended system introduces "
        "a separate mechanism for numeral quantifiers. Person-count numerals are looked "
        "up in a dedicated person-count table, while larger or non-lexicalized "
        "person-count expressions are realized using the maMxi construction. This keeps "
        "count-expression realization separate from ordinary noun plural generation and "
        "allows the input representation to retain roots and grammatical features rather "
        "than precomputed surface forms."
    )

    p2 = (
        "To support numeral quantifiers, the XML input representation was extended in a "
        "controlled manner. A quantifier is represented as an optional element inside "
        "the noun phrase that it modifies, rather than as an independent noun or as a "
        "precomputed surface form. A person-count noun phrase may contain a quantifier "
        "element with type=\"numeral\" and counttype=\"person\", followed by the head "
        "noun with its usual lexical root and grammatical features. Similarly, "
        "object-count expressions use counttype=\"object\". This design preserves the "
        "existing noun phrase structure while allowing the realizer to distinguish "
        "between object-count expressions such as reMdu pustakAlu and person-count "
        "expressions such as ixxaru prakAsAlu. Since the head noun remains represented "
        "by its root and number feature, ordinary noun plural generation continues to "
        "apply to the head noun, while the quantifier is realized by a separate "
        "count-expression module."
    )

    first = insert_after(heading, p1)
    insert_after(first, p2)

    doc.save(PAPER)
    print(f"Updated {PAPER}")
    print(f"Backup {BACKUP}")


if __name__ == "__main__":
    main()
