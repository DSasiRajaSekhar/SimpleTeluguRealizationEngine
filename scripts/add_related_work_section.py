from pathlib import Path
import shutil

from docx import Document


paper_path = Path(r"D:\NLG\Paper.docx")
backup_path = Path(r"D:\NLG\Paper.before_related_work_section.docx")
revised_path = Path(r"D:\NLG\Paper.with_related_work_section.docx")

section_title = "Background and Related Work"

section_paragraphs = [
    (
        "Surface realization has often been studied as a distinct stage in "
        "natural language generation, with SimpleNLG providing an influential "
        "architecture for rule-based English realization (Gatt and Reiter, "
        "2009). Subsequent work on realization for languages other than "
        "English has shown that the general architecture of a modular realizer "
        "must be adapted to the morphological and syntactic properties of each "
        "language. Telugu presents a particularly demanding case because the "
        "surface form of a sentence depends on the interaction of case marking, "
        "agreement, tense-aspect-mode morphology, morphophonemic alternations, "
        "and constituent ordering."
    ),
    (
        "The earlier Telugu surface realization engine addressed this problem "
        "by adapting the SimpleNLG idea to a structured XML input format and a "
        "set of Telugu-specific rule modules (Dokkara, Penumathsa, and "
        "Sripada, 2015). That work established that a symbolic realizer can "
        "generate Telugu sentences when the input provides the required "
        "grammatical features. The present paper builds on that foundation by "
        "using the earlier architecture as the baseline for controlled "
        "experimentation and extension, rather than replacing it with a "
        "statistical or neural generator."
    ),
    (
        "The nominal and verbal components of the system are also connected to "
        "earlier work on Telugu morphology. The noun and pronoun morphology "
        "module is grounded in the treatment of case and number inflection "
        "described in the Telugu noun/pronoun generator (Dokkara, Penumathsa, "
        "and Sripada, 2017). The verb module is grounded in the Telugu verb "
        "morphological generator, which describes verb class identification, "
        "phonetic or morphophonemic alterations, tense-mode suffixes, and "
        "agreement endings (Dokkara, Penumathsa, and Sripada, 2017). These "
        "modules are important because many Telugu realization errors cannot "
        "be corrected by word ordering alone; they require correct inflection "
        "and agreement."
    ),
    (
        "The linguistic decisions in the system are further informed by "
        "descriptive grammar, especially the treatment of Telugu morphology "
        "and verb paradigms in Krishnamurti and Gwynn's A Grammar of Modern "
        "Telugu. This grammar is useful for distinguishing finite tense forms "
        "from verbal adjectives or participial forms, and for interpreting "
        "the behavior of agreement and TAM morphology. In the present work, "
        "such grammatical sources are used not to replace the original system, "
        "but to guide careful extensions and to explain why a particular "
        "baseline output should be treated as a rule-based realization issue."
    ),
]


def has_section(document: Document, title: str) -> bool:
    return any(paragraph.text.strip() == title for paragraph in document.paragraphs)


def main() -> None:
    if not paper_path.exists():
        raise FileNotFoundError(paper_path)

    shutil.copy2(paper_path, backup_path)

    document = Document(str(paper_path))
    if not has_section(document, section_title):
        document.add_heading(section_title, level=1)
        for text in section_paragraphs:
            document.add_paragraph(text)

    try:
        document.save(str(paper_path))
        print(f"Updated: {paper_path}")
    except PermissionError:
        document.save(str(revised_path))
        print(f"Original locked; saved revised copy: {revised_path}")
    print(f"Backup: {backup_path}")


if __name__ == "__main__":
    main()
