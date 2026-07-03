from pathlib import Path
import shutil

from docx import Document


paper_path = Path(r"D:\NLG\Paper.docx")
backup_path = Path(r"D:\NLG\Paper.before_reproducibility_revision.docx")
revised_path = Path(r"D:\NLG\Paper.revised_intro.docx")

revised_paragraphs = [
    (
        "Surface realization is the final stage of natural language generation, "
        "where an abstract linguistic or semantic representation is converted "
        "into a grammatical sentence in a target language. For morphologically "
        "rich languages such as Telugu, this stage is especially challenging "
        "because the realizer must handle case marking, agreement, "
        "tense-aspect-mode morphology, morphophonemic alternations, and "
        "relatively free constituent order. Earlier work on a SimpleNLG-style "
        "surface realization engine for Telugu showed that a rule-based "
        "approach can generate well-formed sentences from structured XML input. "
        "The present work revisits that engine from the perspective of "
        "reproducibility and extensibility: it establishes a documented "
        "baseline, verifies the behavior of the existing realization pipeline, "
        "and identifies the linguistic and engineering points that must be "
        "stabilized before broader extensions can be attempted."
    ),
    (
        "Although the earlier Telugu realizer demonstrated the feasibility of "
        "rule-based surface realization for a morphologically rich Dravidian "
        "language, extending such a system requires a stable and inspectable "
        "baseline. The available implementation contains interacting modules "
        "for XML input processing, phrase construction, noun and pronoun "
        "morphology, verb morphology, agreement, and WX-to-Telugu conversion. "
        "Before adding new linguistic coverage, these modules must be exercised "
        "in a controlled setting so that each XML input passes through the "
        "intended rule-based pipeline and each generated output can be "
        "reproduced. The present work therefore treats baseline stabilization "
        "as a necessary first step: it records the current behavior of the "
        "system, corrects selected baseline errors through rule-based changes, "
        "and preserves a regression set for future extensions."
    ),
    (
        "The contribution of this paper is threefold. First, it reconstructs a "
        "reproducible baseline for the Telugu surface realization engine by "
        "organizing the existing rule-based modules and evaluation inputs into "
        "a controlled experimental setup. Second, it performs a linguistic "
        "audit of the baseline outputs, identifying errors involving agreement "
        "control, XML structure, TAM selection, and irregular verb realization. "
        "Third, it introduces minimal rule-based corrections and records the "
        "resulting outputs as a regression baseline for future extensions. In "
        "doing so, the paper positions the Telugu realizer not as a finished "
        "broad-coverage NLG system, but as an inspectable symbolic baseline "
        "that can support systematic extension and comparison."
    ),
]


def main() -> None:
    if not paper_path.exists():
        raise FileNotFoundError(paper_path)

    shutil.copy2(paper_path, backup_path)

    document = Document(str(paper_path))
    non_empty_paragraphs = [p for p in document.paragraphs if p.text.strip()]

    for index, text in enumerate(revised_paragraphs):
        if index < len(non_empty_paragraphs):
            non_empty_paragraphs[index].text = text
        else:
            document.add_paragraph(text)

    try:
        document.save(str(paper_path))
        print(f"Updated: {paper_path}")
    except PermissionError:
        document.save(str(revised_path))
        print(f"Original locked; saved revised copy: {revised_path}")
    print(f"Backup: {backup_path}")


if __name__ == "__main__":
    main()
