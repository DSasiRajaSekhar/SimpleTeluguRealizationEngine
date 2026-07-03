from pathlib import Path
import shutil

from docx import Document


paper_path = Path(r"D:\NLG\Paper.docx")
backup_path = Path(r"D:\NLG\Paper.before_original_engine_section.docx")
revised_path = Path(r"D:\NLG\Paper.with_original_engine_section.docx")

section_title = "Original Telugu Surface Realization Engine"

section_paragraphs = [
    (
        "The original Telugu surface realization engine was designed as a "
        "rule-based realization system in the spirit of SimpleNLG, but adapted "
        "to the linguistic requirements of Telugu. Its input is a structured "
        "XML representation in which sentence constituents such as subject, "
        "complement, object, modifier, and verb are annotated with grammatical "
        "features. These features include gender, number, person, case marker, "
        "part of speech, and tense-aspect-mode information. The engine uses "
        "this input to build Telugu phrases and sentences through a sequence "
        "of symbolic modules rather than through statistical or neural "
        "generation."
    ),
    (
        "A central strength of the original system is that it treats Telugu "
        "surface realization as a morphology-sensitive problem. Noun and "
        "pronoun forms are generated with case information, while verb forms "
        "are generated through class-based morphophonemic rules, tense-mode "
        "suffixes, and agreement features. This architecture is especially "
        "important for Telugu because the final sentence form depends not only "
        "on word order, but also on case marking, agreement control, TAM "
        "realization, and phonological alternations in the word forms. The "
        "present paper takes this symbolic architecture as its starting point "
        "and uses it as the baseline for controlled extension."
    ),
]


def has_section(document: Document, title: str) -> bool:
    return any(paragraph.text.strip() == title for paragraph in document.paragraphs)


def main() -> None:
    if not paper_path.exists():
        raise FileNotFoundError(paper_path)

    shutil.copy2(paper_path, backup_path)

    document = Document(str(paper_path))
    if not has_section(document, section_title):
        document.add_heading(section_title, level=1)
        for text in section_paragraphs:
            document.add_paragraph(text)

    try:
        document.save(str(paper_path))
        print(f"Updated: {paper_path}")
    except PermissionError:
        document.save(str(revised_path))
        print(f"Original locked; saved revised copy: {revised_path}")
    print(f"Backup: {backup_path}")


if __name__ == "__main__":
    main()
